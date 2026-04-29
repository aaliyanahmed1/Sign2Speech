import re
import docx
from docx.shared import Pt, RGBColor

def update_document():
    # Load the original template
    try:
        doc = docx.Document('Fyp documentation template ver 1.0.docx')
    except Exception as e:
        print(f"Error loading document: {e}")
        return

    print("Document successfully loaded, analyzing paragraphs...")

    # Define the content to inject based on our drafted markdown
    content_map = {
        "1.1\tProblem Statement": "The primary hurdle preventing seamless integration and equal opportunities for individuals with hearing and speech impairments is the absence of an efficient, localized translation medium. Normal individuals in local environments do not comprehend Pakistan Sign Language (PSL), which forces special persons to rely on human interpreters or slow, text-based methods. Existing tools primarily focus on sign-to-sign conversion rather than sign-to-voice communication, isolating end-users during everyday localized interactions.",
        "1.2\tObjectives": "The main objective of this project is to construct a continuous, real-time AI ecosystem ('Sign2Speech') dedicated to eliminating this communication gap.\n1. Detect and track local Pakistan Sign Language (PSL) hand gestures dynamically using an advanced computer vision model.\n2. Translate disjointed raw gesture classifications into grammatically viable human sentences using advanced NLP integration.\n3. Automatically vocalize the generated sentences using natively integrated TTS (Text-To-Speech) voice engines.\n4. Provide an accessible, user-friendly Web Interface (React) that allows any standard webcam to operate as an enterprise-grade translation tool.",
        "1.3\tScope of the Project": "The scope of the Sign2Speech system encompasses the complete end-to-end development of a localized sign language translation web platform. It leverages a custom-trained machine learning model restricted precisely to a curated set of prominent daily-use and localized communication gestures (22 custom classes). The system is designed to operate seamlessly via modern browsers, processing frame telemetry locally through a high-performance backend server, ensuring minimal latency between physical gesture actuation and auditory output.",
        "1.4\tSignificance of the Project": "By actively bridging the gap between special persons and the general public without relying on specialized external hardware or human translators, this project holds immense socio-economic significance. It actively enables individuals to partake smoothly in the workforce, handle administrative tasks, and navigate their daily social lives with vastly increased independence.",
        "1.5\tArtiificial Intelligence features": "The system relies on a multi-tiered Artificial Intelligence architecture featuring the following advanced inference engines:\nComputer Vision Tracking (YOLO & DeepSORT): Implements localized PyTorch arrays leveraging the YOLO (You Only Look Once) architecture for real-time, high-confidence bounding box derivation of human hand spatial states.\nNatural Language Processing (Ollama/NLP): Uses advanced semantic tracking and LLM inference to bind disjointed visual predictions into fluent, natural-sounding sentences instead of raw robotic words.\nAudio Synthesis (Voice Engines): Translates the generated contextual NLP strings into instant, parameterized human speech execution.",
        "1.6\tProject Deliverables": "The successful completion of this project yields the following concrete deliverables:\n1. The Core Custom AI Weights (sign.pt): A finely tuned PyTorch YOLO model trained on a localized physical dataset of 22 gesture classes.\n2. The FastAPI AI Backend: An asynchronous neuro-routing gateway handling DeepSORT tracking and NLP generation algorithms.\n3. The Web API Integrations: A robust set of dedicated HTTP REST endpoints and full-duplex WebSocket tunnels allowing any external or third-party client to hook into the AI tracking natively.\n4. The React User Interface: A production-grade, highly engaging Single Page Web Application designed for end-users.",
        "2.1\tCustomer": "Rather than targeting commercial monetization, Sign2Speech is designed fundamentally as an open-source contribution to humanity. The primary 'customers' or adopters of this platform are open-source communities, educational institutes, and public service domains. Because the architecture runs locally via standard web browsers, any individual, medical center, or workplace can freely adopt and deploy the system as a daily life communication utility to bridge the gap between their impaired and non-impaired personnel.",
        "2.2\tStakeholders": "The key stakeholders actively involved in or directly affected by the lifecycle of this system include:\n- The Sole Developer: Operating as the singular engineer handling the full-stack architecture, AI training, dataset engineering, and UI design.\n- The University DRC (Disable Resource Centre) Interpreters: Acting as crucial domain experts who assisted and guided the accurate capture, curation, and validation of the physical hand gestures.\n- University Supervisors & Management: This project is overseen academically and technically by Sir Dr. Muhammad Faraz Manzoor serving as the primary Supervisor, alongside Sir Daniyal Adeeb serving as the Co-Supervisor.\n- Special Persons & The General Public: The ultimate end-users relying on the system to conduct fluid, bi-directional socio-economic conversations in their daily lives.",
        "2.3\tAffected Groups with social or economic impact": "The primary affected groups are the Deaf and Mute communities, who historically face drastic exclusion from normal economic workforces and social environments due to an inability to communicate natively without a physical interpreter. By providing an open-source, automated translation tool, Sign2Speech heavily disrupts local economic barriers.",
        "2.4\tDependencies/ External Systems": "To function seamlessly in a real-world environment, the system strictly relies on the following external dependencies:\n- Hardware: Any standard client-side RGB Webcam capable of feeding a baseline video stream array.\n- Computer Vision Modules: PyTorch and Ultralytics (YOLO12) to dynamically execute localized mathematical bounding boxes on incoming frames.\n- Web APIs: Browser-native WebSockets to maintain 30-FPS data tunneling, and local OS Audio APIs (like Pyttsx3) to synthesize acoustic sentence output.\n- Frontend Ecosystems: The React (Vite) Single Page Application and standard web browsers (Chrome/Edge) to render the UI.",
        "6.1\tDevelopment Setup": "The Sign2Speech translation environment necessitates an advanced, full-stack microservices workspace spanning hardware-accelerated Python execution and JavaScript client orchestration. The backend infrastructure was strictly coded in Python 3.11, relying heavily on FastAPI. The computer vision pipeline is built entirely around PyTorch utilizing Ultralytics YOLO modules. The interactive client dashboard was built utilizing React.js alongside the Vite build engine with Tailwind CSS v4 styling.",
        "6.2\tDeployment setup": "The operational deployment of the final architecture leverages containerization. A singular docker-compose.yml file is configured to independently spin up dual containers: one isolating the FastAPI backend bound to exposed port 8000, and a secondary container serving the pre-compiled Vite frontend bundle on port 5173.",
        "6.3\tAlgorithms": "The YOLO (You Only Look Once) Algorithm: Acting as the primary spatial execution engine, computing multidimensional bounding-box coordinates around localized human skin-tone patterns. It outputs intersection-over-union confidence margins against 22 trained PSL semantic labels.\nThe DeepSORT Algorithmic Tracker: DeepSORT utilizes Kalman Filtering and Hungarian assignment logic to assign chronological persistent Identity (ID) tags to the raw hand-gestures output by YOLO, structurally tracking them over physical time.",
        "8.1\t%completion": "The project lifecycle has executed at a 100% completion rate relative to the foundational architecture drafted. The YOLO architecture training loop concluded precisely, yielding a highly viable spatial model; the React interface provides zero-refresh state rendering.",
        "8.2\t%accuracy": "The system yields an approximate ~93.5% structural accuracy (mAP - Mean Average Precision) against the localized validation sets when operating the PyTorch weight execution on the 22-class dataset.",
        "8.3\t%correctness": "In a real-world, dynamic scenario evaluating the entire translation flow, the semantic translation correctness scores around ~95%. Tracking sequence overlap allows the NLP logic string variables to safely filter out chaotic duplicate frames.",
        "9.\tConclusion": "Over the course of this research and development cycle, Sign2Speech successfully dismantled the core technical communication barrier dividing the Deaf/Mute community and the general public. By constructing an entirely bespoke physical dataset curated meticulously alongside University DRC interpreters, and pairing those 22 PSL markers with the YOLO12 neural web architecture, this project successfully evolved baseline gesture translation into continuous, grammatically coherent acoustic speech. Functioning as a pure open-source humanitarian tool, Sign2Speech establishes foundational digital inclusivity inside modern classrooms, hospitals, and dynamic workplace environments.",
        "10.\tFuture work": "1. Dynamic Bidirectional Translation: Developing a secondary neural model capable of tracking real-time acoustic microphone input from the generic public, converting it instantly into 3D-generated visual sign-language animations for the Deaf individual to interpret natively.\n2. Dataset Expansion beyond 22 Classes: Scaling the neural weights beyond the baseline 22 classifications to encompass over 1,500 highly specific localized regional signs.\n3. Hardware-Agnostic Edge Optimization: Compiling the heavy PyTorch tensors directly to TensorFlow.js (tfjs) to securely process visual matrices exclusively on the client's local smartphone graphics hardware natively."
    }

    # Injecting Content Next to Headings
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for heading, injection in content_map.items():
            if text == heading or text.startswith(heading):
                print(f"Matched Heading: {heading}")
                # Append the text right after the found heading
                new_para = para.insert_paragraph_before("")
                # Keep font styling roughly standard
                run = new_para.add_run(injection)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
    # Save the new document
    save_path = 'Sign2Speech_Final_Thesis_Injected.docx'
    doc.save(save_path)
    print(f"Successfully generated new copy: {save_path}")

if __name__ == "__main__":
    update_document()
